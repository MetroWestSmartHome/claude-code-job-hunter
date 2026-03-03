#!/usr/bin/env python3
"""
Convert markdown resume to formatted .docx file.

Usage:
    python md-to-docx.py path/to/resume.md

Output:
    Creates resume.docx in the same directory as the input file.

Formatting:
    - Calibri font
    - 10pt body text, 11pt headers
    - Margins: 0.5" top/bottom, 0.7" left/right
    - Section headers: 11pt bold with bottom border
    - Job titles: 11pt bold
    - Job metadata: 10pt (company bold)
    - Job descriptions: 10pt italic
    - Bullets: Word bullet format
"""

import sys
import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_border_to_paragraph(paragraph, border_position='bottom', size=6, color='666666'):
    """Add a border to a paragraph."""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')

    border = OxmlElement(f'w:{border_position}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), str(size))
    border.set(qn('w:space'), '1')
    border.set(qn('w:color'), color)

    pBdr.append(border)
    pPr.append(pBdr)


def parse_markdown_text(text):
    """
    Parse markdown formatting (bold, italic, links) and return list of runs.
    Returns: [(text, bold, italic, hyperlink_url)]
    """
    runs = []

    # Pattern to match **bold**, *italic*, and [text](url)
    # This is a simplified parser - handles basic cases
    pattern = r'(\*\*.*?\*\*|\*.*?\*|\[.*?\]\(.*?\))'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue

        if part.startswith('**') and part.endswith('**'):
            # Bold text
            runs.append((part[2:-2], True, False, None))
        elif part.startswith('*') and part.endswith('*'):
            # Italic text
            runs.append((part[1:-1], False, True, None))
        elif part.startswith('[') and '](' in part:
            # Hyperlink
            match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if match:
                link_text, url = match.groups()
                runs.append((link_text, False, False, url))
        else:
            # Plain text
            runs.append((part, False, False, None))

    return runs


def add_formatted_text(paragraph, text, default_size=20, default_bold=False, default_italic=False):
    """Add text to paragraph with markdown formatting."""
    runs_data = parse_markdown_text(text)

    for text_content, is_bold, is_italic, url in runs_data:
        if url:
            # Add hyperlink
            add_hyperlink(paragraph, url, text_content, default_size)
        else:
            run = paragraph.add_run(text_content)
            run.font.name = 'Calibri'
            run.font.size = Pt(default_size / 2)  # Convert half-points to points
            run.font.bold = is_bold or default_bold
            run.font.italic = is_italic or default_italic


def add_hyperlink(paragraph, url, text, font_size=20):
    """Add a hyperlink to a paragraph."""
    # This is a workaround - python-docx doesn't have built-in hyperlink support
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Style for hyperlink
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    # Set font properties
    run = paragraph.runs[-1]
    run.font.name = 'Calibri'
    run.font.size = Pt(font_size / 2)


def convert_resume(md_file):
    """Convert markdown resume to .docx."""

    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Create document
    doc = Document()

    # Set margins: 0.5" top/bottom, 0.7" left/right
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Process lines
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # H1 (Name) - # [YOUR_NAME]
        if line.startswith('# ') and not line.startswith('## '):
            name = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(name)
            run.font.name = 'Calibri'
            run.font.size = Pt(16)
            run.font.bold = True

        # H2 (Section headers) - ## PROFESSIONAL SUMMARY
        elif line.startswith('## '):
            header = line[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(9)
            p.paragraph_format.space_after = Pt(4.5)
            add_border_to_paragraph(p, 'bottom', 4, '666666')
            run = p.add_run(header)
            run.font.name = 'Calibri'
            run.font.size = Pt(11.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(51, 51, 51)

        # H3 (Job titles) - ### [Your Job Title]
        elif line.startswith('### '):
            job_title = line[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4.5)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(job_title)
            run.font.name = 'Calibri'
            run.font.size = Pt(11.5)
            run.font.bold = True

        # Horizontal rule (---)
        elif line.strip() == '---':
            # Add spacing
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)

        # Bold text starting line (contact info, job metadata, competencies)
        elif line.startswith('**') and '**' in line[2:]:
            # Could be contact info, job metadata, or competency labels

            # Check if it's contact info (multiple ** pairs on same line)
            if line.count('**') >= 4:
                # Contact info line
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(6)
                add_formatted_text(p, line, default_size=19)

            # Check if next line exists and starts with ** (competency format)
            elif i + 1 < len(lines) and not lines[i + 1].startswith('**') and not lines[i + 1].startswith('-'):
                # Single line bold (likely job metadata or competency label)
                # Job metadata format: **Company** | Location | Dates
                if '|' in line:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(6)
                    add_formatted_text(p, line, default_size=20)
                else:
                    # Competency label: **Category:** items
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(6)
                    add_formatted_text(p, line, default_size=20)
            else:
                # Regular paragraph with bold
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                add_formatted_text(p, line, default_size=20)

        # Italic text (job description)
        elif line.startswith('*') and not line.startswith('**') and not line.startswith('- '):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(7.5)
            add_formatted_text(p, line, default_size=20, default_italic=True)

        # Bullet points
        elif line.startswith('- '):
            bullet_text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4.5)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            add_formatted_text(p, bullet_text, default_size=20)

        # Regular paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            add_formatted_text(p, line, default_size=20)

        i += 1

    # Save document
    output_file = Path(md_file).with_suffix('.docx')
    doc.save(output_file)
    print(f"Generated: {output_file}")
    return output_file


def main():
    if len(sys.argv) < 2:
        print("Usage: python md-to-docx.py path/to/resume.md")
        sys.exit(1)

    md_file = sys.argv[1]

    if not os.path.exists(md_file):
        print(f"Error: File not found: {md_file}")
        sys.exit(1)

    if not md_file.endswith('.md'):
        print(f"Error: Input file must be a .md file")
        sys.exit(1)

    convert_resume(md_file)


if __name__ == '__main__':
    main()
