#!/usr/bin/env python3
"""
Generate cover letter .docx from template markdown.

Usage:
    python cover-letter-template.py --input cover-letter-template.md

    OR with inline content:

    python cover-letter-template.py \
        --company "CompanyName" \
        --role "Job Title" \
        --why-interested "..." \
        --paragraph1 "..." \
        --paragraph2 "..." \
        --paragraph3 "..." \
        [--paragraph4 "..."]

Output:
    Creates cover-letter.docx in the same directory as the input file.

Formatting:
    - Calibri 11pt font
    - 1" margins all around
    - Professional business letter format
    - Proper spacing between paragraphs
"""

import sys
import argparse
import re
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def parse_template_file(file_path):
    """Parse cover letter template markdown file."""
    content = Path(file_path).read_text(encoding='utf-8')

    # Extract sections
    sections = {}

    # Extract company and role from header
    company_match = re.search(r'# Cover Letter Template - (.+)', content)
    role_match = re.search(r'\*\*Role:\*\* (.+)', content)

    sections['company'] = company_match.group(1).strip() if company_match else "Company"
    sections['role'] = role_match.group(1).strip() if role_match else "Role"

    # Extract content sections
    why_match = re.search(r'### Why Interested\n(.+?)(?=\n###|\n---|\Z)', content, re.DOTALL)
    p1_match = re.search(r'### Paragraph 1: Experience Alignment\n(.+?)(?=\n###|\n---|\Z)', content, re.DOTALL)
    p2_match = re.search(r'### Paragraph 2: Technical Depth\n(.+?)(?=\n###|\n---|\Z)', content, re.DOTALL)
    p3_match = re.search(r'### Paragraph 3: Leadership & Approach\n(.+?)(?=\n###|\n---|\Z)', content, re.DOTALL)
    p4_match = re.search(r'### Paragraph 4: Gap Addressing.*?\n(.+?)(?=\n###|\n---|\Z)', content, re.DOTALL)

    sections['why_interested'] = why_match.group(1).strip() if why_match else None
    sections['paragraph1'] = p1_match.group(1).strip() if p1_match else None
    sections['paragraph2'] = p2_match.group(1).strip() if p2_match else None
    sections['paragraph3'] = p3_match.group(1).strip() if p3_match else None
    sections['paragraph4'] = p4_match.group(1).strip() if p4_match else None

    # Remove template instructions (lines starting with [)
    for key in ['why_interested', 'paragraph1', 'paragraph2', 'paragraph3', 'paragraph4']:
        if sections[key]:
            # Remove lines that are just template instructions
            lines = [line for line in sections[key].split('\n') if not line.strip().startswith('[')]
            sections[key] = '\n'.join(lines).strip()

            # If empty after removing instructions, set to None
            if not sections[key]:
                sections[key] = None

    return sections


def create_cover_letter_docx(sections, output_path):
    """Create formatted cover letter .docx."""
    doc = Document()

    # Set page margins (1" all around)
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Default style: Calibri 11pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Contact information (right-aligned)
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # TODO: Replace with your information or read from config
    contact.add_run("[YOUR_NAME]\n").bold = True
    contact.add_run("[Your City, State ZIP]\n")
    contact.add_run("[Your Phone]\n")
    contact.add_run("[your.email@example.com]\n")
    contact.add_run("linkedin.com/in/yourprofile")

    # Add spacing
    doc.add_paragraph()

    # Date
    date_para = doc.add_paragraph(datetime.now().strftime('%B %d, %Y'))

    # Add spacing
    doc.add_paragraph()

    # Greeting
    greeting = doc.add_paragraph(f"Dear {sections['company']} Hiring Team,")

    # Add spacing
    doc.add_paragraph()

    # Opening paragraph (why interested)
    if sections['why_interested']:
        opening = doc.add_paragraph(sections['why_interested'])

    # Body paragraphs
    for i in range(1, 5):
        key = f'paragraph{i}'
        if sections.get(key):
            doc.add_paragraph()  # Spacing
            para = doc.add_paragraph(sections[key])

    # Add spacing before closing
    doc.add_paragraph()

    # Closing
    closing = doc.add_paragraph("Thank you for your consideration. I look forward to discussing how my experience aligns with your needs.")

    # Add spacing
    doc.add_paragraph()

    # Signature
    signature = doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph()
    sig_name = doc.add_paragraph("[YOUR_NAME]")  # TODO: Replace with your name

    # Save
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description='Generate cover letter .docx from template')

    # Option 1: Use template file
    parser.add_argument('--input', help='Input markdown template file')

    # Option 2: Inline content
    parser.add_argument('--company', help='Company name')
    parser.add_argument('--role', help='Job title')
    parser.add_argument('--why-interested', help='Why interested paragraph')
    parser.add_argument('--paragraph1', help='Experience alignment paragraph')
    parser.add_argument('--paragraph2', help='Technical depth paragraph')
    parser.add_argument('--paragraph3', help='Leadership approach paragraph')
    parser.add_argument('--paragraph4', help='Gap addressing paragraph (optional)')

    parser.add_argument('--output', help='Output .docx file path (default: cover-letter.docx in same dir as input)')

    args = parser.parse_args()

    # Determine source: template file or inline
    if args.input:
        # Parse template file
        input_path = Path(args.input)
        if not input_path.exists():
            print(f"Error: Input file not found: {input_path}")
            return 1

        sections = parse_template_file(input_path)

        # Validate sections are filled in
        missing = []
        for key in ['why_interested', 'paragraph1', 'paragraph2', 'paragraph3']:
            if not sections.get(key):
                missing.append(key)

        if missing:
            print(f"Error: Template not filled in. Missing sections: {', '.join(missing)}")
            print(f"\nPlease edit {input_path} and fill in the content sections.")
            return 1

        # Default output path
        if not args.output:
            output_path = input_path.parent / "cover-letter.docx"
        else:
            output_path = Path(args.output)

    elif args.company and args.role and args.why_interested and args.paragraph1 and args.paragraph2 and args.paragraph3:
        # Inline content
        sections = {
            'company': args.company,
            'role': args.role,
            'why_interested': args.why_interested,
            'paragraph1': args.paragraph1,
            'paragraph2': args.paragraph2,
            'paragraph3': args.paragraph3,
            'paragraph4': args.paragraph4,
        }

        # Default output path
        if not args.output:
            output_path = Path("cover-letter.docx")
        else:
            output_path = Path(args.output)

    else:
        print("Error: Must provide either --input or all inline content arguments")
        print("\nUsage 1 (template file):")
        print("  python cover-letter-template.py --input cover-letter-template.md")
        print("\nUsage 2 (inline content):")
        print("  python cover-letter-template.py --company 'CompanyName' --role 'Job Title' \\")
        print("    --why-interested '...' --paragraph1 '...' --paragraph2 '...' --paragraph3 '...'")
        return 1

    # Create cover letter
    print(f"Generating cover letter: {output_path}")
    create_cover_letter_docx(sections, output_path)
    print(f"✅ Cover letter created: {output_path}")

    return 0


if __name__ == "__main__":
    exit(main())
